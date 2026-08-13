# -*- coding: utf-8 -*-
"""修改后测问卷 docx：
1) 修正 Q9 双重否定表述（保持正相计分语义）
2) 新增第四部分：自我评估（效标效度必需：整体自评0-100 + 三维度自评 + 生态效度）
3) 新增第五部分：开放式反馈（最难环节 / 改进建议）
"""
import copy, sys
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx.shared import Pt

SRC = 'AI辅助工作实验后测-8.13.docx'
d = docx.Document(SRC)

# ---------- 1) 修改 Q9（段落下标 20）----------
p9 = d.paragraphs[20]
assert '过度怀疑' in p9.text, f'Q9定位失败: {p9.text}'
# runs: 0='9. ' 1=题干 2='AI' 3=尾巴 4..7='[单选题] *'
p9.runs[0].text = '9. '
p9.runs[1].text = '当确认AI内容无误后，我能及时停止核查，不让过度怀疑拖慢效率。'
p9.runs[2].text = ''
p9.runs[3].text = ''

# ---------- 工具函数 ----------
def add_para(text, bold=False, italic=False):
    p = d.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_scale_table():
    """深拷贝现有 1x7 量表表（很不符合|○1-5|很符合），保证格式完全一致"""
    tbl = copy.deepcopy(d.tables[0]._tbl)
    d.element.body.append(tbl)

# ---------- 2) 第四部分 自我评估 ----------
add_para('')
add_para('---第3页---', bold=True)
add_para('')
add_para('第四部分 自我评估', bold=True)
add_para('以下问题请根据你刚才在测验中的真实表现作答。')

add_para('16. 请为本次测验中自己的整体表现打分：____（0-100 分） [填空题] *')
add_para('')

add_para('17. 我能准确判断AI结果何时可信、何时该怀疑。 [单选题] *')
add_scale_table()
add_para('18. 我能主动核验AI输出中的关键数据、事实与条款。 [单选题] *')
add_scale_table()
add_para('19. 面对可能违规或越权的内容，我能守住边界、不越线操作。 [单选题] *')
add_scale_table()
add_para('20. 本次测验中我的表现，能反映我平时在真实工作中的行为方式。 [单选题] *')
add_scale_table()

# ---------- 3) 第五部分 开放式反馈 ----------
add_para('')
add_para('第五部分 反馈与建议', bold=True)
add_para('21. 这次测验中，你觉得最难或印象最深的环节是什么？为什么？ [问答题]')
add_para('')
add_para('22. 你对这套测验系统有什么改进建议？ [问答题]')

d.save(SRC)
print('已保存:', SRC)

# ---------- 校验 ----------
d2 = docx.Document(SRC)
txt = [p.text for p in d2.paragraphs]
qs = [t for t in txt if t.strip() and t.strip()[0].isdigit() and '. ' in t]
print('题项总数:', len(qs), '| 表格数:', len(d2.tables))
for t in qs:
    print('  ', t[:50])
